import streamlit as st
from google import genai
from google.genai import types
import PIL.Image
import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Set up the web app
st.set_page_config(page_title="Master Court Assistant", page_icon="⚖️", layout="wide")
st.title("⚖️ Master Court Assistant")

try:
    api_key = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("API Key not found. Please check Streamlit secrets.")
    st.stop()

client = genai.Client(api_key=api_key)

# ==========================================
# FUNCTION 1: COURT DAILY REPORT FORMAT
# ==========================================
def add_daily_report_to_word(doc, data, mohrir_name):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

    header_table = doc.add_table(rows=2, cols=3)
    po_val = str(data.get("po_name", ""))
    po_text = po_val if po_val and po_val != "-" and po_val.lower() != "none" else "............................"
    abhi_val = str(data.get("abhiyojak_name", ""))
    abhi_text = abhi_val if abhi_val and abhi_val != "-" and abhi_val.lower() != "none" else "श्री संजीव सिंह"
    date_val = str(data.get("report_date", ""))
    date_text = date_val if date_val and date_val != "-" and date_val.lower() != "none" else "............................."
    
    p_left = header_table.cell(0, 0).paragraphs[0]
    p_left.add_run(f"PO- श्री {po_text}\nआरोप बनने का दि० ....................").bold = True
    p_center = header_table.cell(0, 1).paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.add_run("न्यायालय डेली रिपोर्ट").bold = True
    p_center.runs[0].font.size = Pt(16)
    p_right = header_table.cell(0, 2).paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run(f"अभियोजक का नाम- {abhi_text}.\nदिनांक - {date_text}").bold = True
    
    court_cell = header_table.cell(1, 0)
    court_cell.merge(header_table.cell(1, 2))
    court_cell.paragraphs[0].add_run("न्यायालय - ACJM - Kadipur सुल्तानपुर").bold = True
    court_cell.paragraphs[0].runs[0].font.size = Pt(14)
    doc.add_paragraph("")
    
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    col_widths = [Inches(1.2), Inches(2.0), Inches(3.2), Inches(3.6)]
    labels = ['थाना', 'अ०सं०', 'वाद सं०', 'धारा', 'वादी का नाम व पता', 'विवेचक का नाम', 'निर्णय का दि०', 'अभियुक्त का नाम व पता']
    for i, label in enumerate(labels):
        table.cell(i, 0).text = label
        
    table.cell(0, 1).text = str(data.get("thana", "-"))
    table.cell(1, 1).text = str(data.get("apr_sankhya", "-"))
    table.cell(2, 1).text = str(data.get("vaad_sankhya", "-"))
    table.cell(3, 1).text = str(data.get("dhara", "-"))
    table.cell(4, 1).text = str(data.get("vaadi", "-"))
    table.cell(5, 1).text = str(data.get("vivechak", "-"))
    table.cell(6, 1).text = str(data.get("nirnay_date", "-"))
    table.cell(7, 1).text = str(data.get("abhiyukt", "-"))
    
    ghatna_val = str(data.get("ghatna", ""))
    dhara_val = str(data.get("dhara", ""))
    if not ghatna_val or ghatna_val == "-" or ghatna_val.lower() == "none":
        if "आर्म्स" in dhara_val or "Arms" in dhara_val:
            ghatna_val = f"अभियुक्त के पास से अवैध शस्त्र बरामद हुआ। अपराध धारा {dhara_val} के अंतर्गत दण्डनीय है।"
        elif "323" in dhara_val:
            ghatna_val = f"अभियुक्तगण ने वादी के साथ गाली-गलौज की तथा मारपीट कर साधारण उपहति (चोट) कारित की। अपराध धारा {dhara_val} के अंतर्गत दण्डनीय है।"
        else:
            ghatna_val = f"अभियुक्त के विरुद्ध मामला पंजीकृत किया गया। अपराध धारा {dhara_val} के अंतर्गत दण्डनीय है।"

    table.cell(0, 2).text = "घटना का संक्षिप्त विवरण"
    table.cell(0, 2).paragraphs[0].runs[0].bold = True
    c_ghatna = table.cell(1, 2)
    c_ghatna.merge(table.cell(7, 2))
    c_ghatna.text = "" 
    c_ghatna.paragraphs[0].add_run(ghatna_val).font.color.rgb = RGBColor(0, 0, 255) 
    
    table.cell(0, 3).text = "न्यायालय के आदेश का विवरण"
    table.cell(0, 3).paragraphs[0].runs[0].bold = True
    c_adesh = table.cell(1, 3)
    c_adesh.merge(table.cell(7, 3))
    c_adesh.text = "" 
    c_adesh.paragraphs[0].add_run(str(data.get("adesh", "-"))).font.color.rgb = RGBColor(255, 0, 0) 
    
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    doc.add_paragraph("\n")
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.cell(0, 0).paragraphs[0].add_run(f"{mohrir_name}\nनाम व हस्ताक्षर कोर्ट मोहर्रिर").bold = True
    sig_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_table.cell(0, 1).paragraphs[0].add_run("हस्ताक्षर अभियोजक").bold = True
    doc.add_page_break()

# ==========================================
# FUNCTION 2: SOOCHNA (NOTICE) FORMAT
# ==========================================
def add_soochna_to_word(doc, data, action_type, mohrir_name):
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_head.add_run("सूचना\n")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(18)
    
    court_name = str(data.get("court_name", "JM-JD"))
    run_court = p_head.add_run(f"न्यायालय {court_name}\nSULTANPUR")
    run_court.bold = True
    run_court.font.size = Pt(14)
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(4.0)
    
    accused = str(data.get("accused", "...................................."))
    table.cell(0, 0).text = f"नाम पता अभियुक्त-\n{accused}"
    
    crime = str(data.get("crime_no", "............."))
    dhara = str(data.get("dhara", "............."))
    thana = str(data.get("thana", "............."))
    table.cell(0, 1).text = f"मु ०अ ०सं०- {crime}\nधारा - {dhara}\n\nथाना- {thana}\nजिला- सुल्तानपुर"
    
    thana_sho = str(data.get("thana", "............."))
    date_app = str(data.get("date", ".................."))
    
    p_body = doc.add_paragraph()
    p_body.add_run(f"\nसेवा में,\n      श्रीमान SHO\n      थाना {thana_sho}, जिला सुल्तानपुर\nमहोदय,\n")
    action_text = "जमानत" if action_type == "जमानत (Bail)" else "NBW RECALL"
    p_body.add_run(f"      निवेदन के साथ अवगत कराना है कि मु ०अ ०सं० व धारा उपरोक्त में अभि० उपरोक्त दिनांक {date_app} को न्यायालय के समक्ष उपस्थित होकर {action_text} करा लिया है।\n{action_text} सूचना आवश्यक कार्यवाही हेतु प्रेषित है।")
    
    p_foot = doc.add_paragraph(f"\n\n{mohrir_name}\nको०मो०")
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_foot.runs:
        run.bold = True
    doc.add_page_break()

# ==========================================
# FUNCTION 3: WITNESS/SUMMONS REPORT FORMAT
# ==========================================
def add_witness_report_to_word(doc, data, mohrir_name):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    court_name = str(data.get("court_name", "JM (JD) OUTLINE COURT कादीपुर सुल्तानपुर"))
    p_head.add_run(f"न्यायालय {court_name}\n").bold = True
    p_head.runs[0].font.size = Pt(16)
    
    report_date = str(data.get("report_date", ".................."))
    p_head.add_run(f"साक्षिगणों को निर्गत समन व उनकी उपस्थिति के संबंध में विवरण दिनांक {report_date}").font.size = Pt(14)
    
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    headers = [
        "जनपद", "साक्ष्य हेतु\nनियत वाद", "निर्गत कुल\nसम्मन", 
        "तामीला\nहोकर\nप्राप्त\nसम्मन", "निर्गत कुल\nवारंट", 
        "तामीला\nहोकर प्राप्त\nवारंट", "उपस्थित आए\nसाक्षियों की\nसंख्या", 
        "परीक्षित\nसाक्षियों की\nसंख्या"
    ]
    for i, head_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = head_text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        
    keys = ["janpad", "niyat_vaad", "nirgat_samman", "tamila_samman", "nirgat_warrant", "tamila_warrant", "upasthit_sakshi", "parikshit_sakshi"]
    for i, key in enumerate(keys):
        cell = table.cell(1, i)
        cell.text = str(data.get(key, "00"))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("")
    vc_count = str(data.get("vc_count", "0"))
    p_foot = doc.add_paragraph(f"VC से होने वाली गवाही सूचना - {vc_count}")
    p_foot.runs[0].font.size = Pt(14)
    doc.add_page_break()

# ==========================================
# FUNCTION 4: TOP 10 REPORT FORMAT
# ==========================================
def add_top_10_report_to_word(doc, rows_data):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.add_run("न्यायालय ACJM Kadipur सुल्तानपुर\n").bold = True
    p_head.runs[0].font.size = Pt(16)
    p_head.add_run("टॉप 10 की सूचना").bold = True
    p_head.runs[1].font.size = Pt(14)
    doc.add_paragraph("")

    table = doc.add_table(rows=len(rows_data) + 1, cols=10)
    table.style = 'Table Grid'
    table.autofit = False

    headers = [
        "अभियुक्त का नाम", "मु०अ०सं०", "धारा", "थाना", "वाद संख्या", 
        "कुल गवाहों\nकी संख्या", "परीक्षित\nगवाहों", "परीक्षित हेतु\nशेष गवाहों की संख्या", 
        "वर्तमान स्थिति", "अग्रिम तिथि"
    ]
    col_widths = [Inches(1.5), Inches(0.9), Inches(1.1), Inches(0.8), Inches(0.9), Inches(0.8), Inches(0.8), Inches(1.0), Inches(1.0), Inches(0.8)]

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    for r_idx, row_vals in enumerate(rows_data):
        for c_idx, val in enumerate(row_vals):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_page_break()

# ==========================================
# FUNCTION 5: ADVOCATE CASES REPORT FORMAT
# ==========================================
def add_advocate_cases_report_to_word(doc, rows_data):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.add_run("न्यायालय अपर मुख्य न्यायिक मजिस्ट्रेट, कादीपुर जिला- सुल्तानपुर\n").bold = True
    p_head.runs[0].font.size = Pt(15)
    p_head.add_run('"अधिवक्ता के विरुद्ध दर्ज अभियोगों के सम्बन्ध में सूचना"').bold = True
    p_head.runs[1].font.size = Pt(13)
    doc.add_paragraph("")

    table = doc.add_table(rows=len(rows_data) + 1, cols=10)
    table.style = 'Table Grid'
    table.autofit = False

    headers = [
        "क्र० सं०", "मु०अ०सं०", "धारा", "जिला", "थाना", 
        "अभियुक्त व पिता\nका नाम", "न्यायालय\nका नाम", "न्यायालय में\nवाद की स्थिति", 
        "नियत तिथि", "नियत तिथि पर\nहुई अग्रिम कार्यवाही"
    ]
    col_widths = [Inches(0.6), Inches(0.9), Inches(1.1), Inches(0.8), Inches(0.8), Inches(1.5), Inches(1.2), Inches(0.9), Inches(0.9), Inches(1.3)]

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    for r_idx, row_vals in enumerate(rows_data):
        for c_idx, val in enumerate(row_vals):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_page_break()

# ==========================================
# FUNCTION 6: REMAND SHEET FORMAT (NEW)
# ==========================================
def add_remand_sheet_to_word(doc, data):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11.0)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1.0)

    court_name = str(data.get("court_name", "अपर मुख्य न्यायिक मजिस्ट्रेट-SD कादीपुर सुलतानपुर"))
    
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.add_run("रिमाण्ड शीट\n").bold = True
    p_head.runs[0].font.size = Pt(16)
    p_head.add_run(f"न्यायालय {court_name}").bold = True
    p_head.runs[1].font.size = Pt(13)
    doc.add_paragraph("")

    crime_no = str(data.get("crime_no", ".................."))
    dhara = str(data.get("dhara", ".................."))
    thana = str(data.get("thana", ".................."))
    janpad = str(data.get("janpad", "सुल्तानपुर"))
    
    p_details = doc.add_paragraph()
    p_details.add_run(f"मु०अ०सं०- {crime_no}\nधारा- {dhara}\nथाना- {thana}\nजनपद- {janpad}\n\nराज्य बनाम\n\n1\n\n")
    
    arrest_date = str(data.get("arrest_date", ".................."))
    accused_name = str(data.get("accused_name", "उपरोक्त अभियुक्त"))
    
    p_body = doc.add_paragraph()
    p_body.add_run(f"आज दिनांक-{arrest_date} को उक्त अभियोग में उपरोक्त अभियुक्त ({accused_name}) को न्यायालय के समक्ष पुलिस द्वारा गिरफ्तार कर प्रस्तुत किया गया, जिसकी न्यायिक अभिरक्षा में रिमाण्ड किये जाने हेतु विवेचना अधिकारी द्वारा प्रार्थना की गयी। मैंने प्रथम सूचना रिपोर्ट व केस डायरी का अवलोकन किया जिसके अवलोकन से मेरा मत है कि अभियुक्त उपरोक्त को धारा-187 BNSS (भारतीय नागरिक सुरक्षा संहिता) के अंतर्गत न्यायिक अभिरक्षा में रिमाण्ड किये जाने का आधार पर्याप्त है।\n\n")
    
    p_adesh_head = doc.add_paragraph()
    p_adesh_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_adesh_head.add_run("आदेश").bold = True
    p_adesh_head.runs[0].font.size = Pt(14)
    
    remand_from = str(data.get("remand_from", ".................."))
    remand_to = str(data.get("remand_to", ".................."))
    jail_date = str(data.get("jail_date", ".................."))
    
    p_adesh_body = doc.add_paragraph()
    p_adesh_body.add_run(f"अभियुक्त उपरोक्त का दिनांक-{remand_from} से दिनांक-{remand_to} तक न्यायिक अभिरक्षा में रिमाण्ड स्वीकृत किया जाता है। अभियुक्त को जिला कारागार से तलब कर दिनांक-{jail_date} को मेरे समक्ष प्रस्तुत किया जाए।\n\n\n")
    
    po_name = str(data.get("po_name", "अपर मुख्य न्यायिक मजिस्ट्रेट-SD\nकादीपुर, सुलतानपुर।"))
    date_sign = str(data.get("date_sign", ".................."))
    
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_foot.add_run(f"{po_name}\nदिनांक-{date_sign}")
    doc.add_page_break()

# ==========================================
# FUNCTION 7: COGNIZANCE / SANJAN ORDER FORMAT (NEW)
# ==========================================
def add_sanjan_order_to_word(doc, data):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11.0)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1.0)

    court_name = str(data.get("court_name", "अपर मुख्य न्यायिक मजिस्ट्रेट, कादीपुर, सुलतानपुर"))
    
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.add_run("न्यायालय\n").bold = True
    p_head.runs[0].font.size = Pt(16)
    p_head.add_run(f"{court_name}").bold = True
    p_head.runs[1].font.size = Pt(14)
    doc.add_paragraph("")

    upasthit = str(data.get("upasthit", "सहायक अभियोजन अधिकारी / कोर्ट मोहर्रिर"))
    vaad_no = str(data.get("vaad_no", "............."))
    comp_no = str(data.get("comp_no", "............."))
    cr_no = str(data.get("cr_no", "............."))
    chargesheet_no = str(data.get("chargesheet_no", "............."))
    crime_no = str(data.get("crime_no", "............."))
    dhara = str(data.get("dhara", "............."))
    thana = str(data.get("thana", "............."))
    janpad = str(data.get("janpad", "सुल्तानपुर"))
    order_date = str(data.get("order_date", "............."))
    
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"उपस्थित: {upasthit}\nवाद सं०- {vaad_no}\nकम्प्यूटर फाइलिंग सं०- {comp_no}\nसी०एन०आर०सं०- {cr_no}\nआरोप पत्र सं०- {chargesheet_no}\nमुकदमा अपराध संख्या- {crime_no}\nअन्तर्गत धारा- {dhara}\nथाना- {thana}\nजनपद- {janpad}\nदिनांक- {order_date}\n\n")
    
    accused_list = str(data.get("accused_list", "1...................................................."))
    initial_dhara = str(data.get("initial_dhara", "...................................................."))
    cognizance_dhara = str(data.get("cognizance_dhara", "...................................................."))
    next_date = str(data.get("next_date", "...................................................."))
    po_name = str(data.get("po_name", "(विश्वजीत सिंह)\nअपर मुख्य न्यायिक मजिस्ट्रेट,\nकादीपुर जनपद सुल्तानपुर।"))
    
    p_body = doc.add_paragraph()
    p_body.add_run(f"आज पत्रावली संज्ञान के बिंदु पर सुनवाई हेतु प्रस्तुत हुई। संज्ञान के बिंदु पर पत्रावली/अभिलेक्स का अवलोकन किया गया।\nअवलोकन से विदित होता है कि थाना {thana}, जनपद {janpad} के उक्त प्रकरण में विवेचक द्वारा अभियुक्त / अभियुक्तगण—\n{accused_list}\n\nके विरुद्ध धारा {initial_dhara} के अन्तर्गत आरोप पत्र न्यायालय में प्रस्तुत किया गया है। केस डायरी एवं संलग्न प्रपत्रों का अवलोकन किया गया। पत्रावली पर उपलब्ध साक्षियों तथा गवाहों के कथनों के आधार पर उपरोक्त अभियुक्तगण के विरुद्ध प्रथमदृष्ट्या धारा {cognizance_dhara} के अन्तर्गत संज्ञेय अपराध कारित होना प्रतीत होता है।\nअतः, उपरोक्त अभियुक्त/अभियुक्तगण—\n{accused_list}\n\nके विरुद्ध धारा {cognizance_dhara} के अन्तर्गत अपराध का संज्ञान लिया जाता है।\nपत्रावली नियमित वाद के रूप में पंजीकृत की जाए तथा अभियुक्तगण को समन/वारंट (या कारागार से) तलब किया जाए।\nवास्ते उपस्थिति दिनांक— {next_date} नियत की जाती है।\n\n\n\n")
    
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_foot.add_run(po_name)
    doc.add_page_break()


# ==========================================
# MAIN APP UI
# ==========================================
st.markdown("### ⚙️ Global Settings")
global_mohrir = st.text_input("✍️ Court Mohrir Name (कोर्ट मोहर्रिर):", value="का ० अभय राज सिंह")
st.markdown("---")

doc_type = st.selectbox("📑 Select Document Type to Generate:", [
    "Court Daily Report (डेली रिपोर्ट)", 
    "Soochna (सूचना - Jamanat / NBW Recall)",
    "Witness/Summons Report (साक्षी/समन विवरण)",
    "Top 10 Report (टॉप 10 की सूचना)",                      
    "Cases Against Advocates Report (अधिवक्ता के विरुद्ध अभियोग)",
    "Remand Sheet (रिमाण्ड शीट)",                     # NEW
    "Cognizance / Sanjan Order (संज्ञान आदेश)"        # NEW
])
st.markdown("---")

# ----------------- UI: DAILY REPORT -----------------
if doc_type == "Court Daily Report (डेली रिपोर्ट)":
    mode = st.radio("How should multiple photos be handled?", ["📂 Combine all photos into ONE report", "📄 Process SEPARATE reports"])
    uploaded_files = st.file_uploader("Upload Photos (JPG/PNG)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    if uploaded_files:
        preview_cols = st.columns(len(uploaded_files) if len(uploaded_files) < 4 else 4)
        for i, file in enumerate(uploaded_files):
            preview_cols[i % 4].image(file, caption=f"Page {i+1}", use_container_width=True)
            
    audio_file = st.audio_input("🎙️ Record Dictation (Overrides Photo)")
    
    with st.expander("📝 Manual Text Overrides"):
        head_col1, head_col2, head_col3 = st.columns(3)
        man_po = head_col1.text_input("PO- श्री (Judge):")
        man_abhiyojak = head_col2.text_input("अभियोजक (Prosecutor):")
        man_date = head_col3.text_input("दिनांक (Date):")
        man_thana = st.text_input("थाना (Thana):")
        col1, col2 = st.columns(2)
        man_apr = col1.text_input("अ०सं० (Crime No):")
        man_dhara = col1.text_input("धारा (Sections):")
        man_vaadi = col1.text_area("वादी का नाम (Complainant):", height=68)
        man_nirnay = col1.text_input("निर्णय का दि०:")
        man_vaad = col2.text_input("वाद सं० (Case No):")
        man_vivechak = col2.text_input("विवेचक का नाम:")
        man_abhiyukt = col2.text_area("अभियुक्त का नाम (Accused):", height=68)
        man_ghatna = st.text_area("घटना का संक्षिप्त विवरण (Ghatna):")
        man_adesh = st.text_area("न्यायालय के आदेश का विवरण (Aadesh):")

    if (uploaded_files or audio_file) and st.button("Generate Daily Report", type="primary"):
        with st.spinner("Processing..."):
            doc = Document()
            prompt = f"""
            Extract information into JSON with keys: "po_name", "abhiyojak_name", "report_date", "thana", "apr_sankhya", "vaad_sankhya", "dhara", "vaadi", "vivechak", "nirnay_date", "abhiyukt", "ghatna", "adesh".
            OVERRIDES: PO:{man_po} Abhiyojak:{man_abhiyojak} Date:{man_date} Thana:{man_thana} Crime:{man_apr} Case:{man_vaad} Dhara:{man_dhara} Vaadi:{man_vaadi} Vivechak:{man_vivechak} NirnayDate:{man_nirnay} Accused:{man_abhiyukt} Ghatna:{man_ghatna} Adesh:{man_adesh}. Output ONLY JSON.
            """
            audio_part = types.Part.from_bytes(data=audio_file.getvalue(), mime_type=audio_file.type) if audio_file else None
            try:
                if "Combine" in mode and uploaded_files:
                    contents = [PIL.Image.open(f) for f in uploaded_files]
                    if audio_part: contents.append(audio_part)
                    contents.append(prompt)
                    res = client.models.generate_content(model="gemini-3.6-flash", contents=contents).text.strip()
                    add_daily_report_to_word(doc, json.loads(res.replace("```json","").replace("```","").strip()), global_mohrir)
                elif uploaded_files:
                    for f in uploaded_files:
                        contents = [PIL.Image.open(f), audio_part, prompt] if audio_part else [PIL.Image.open(f), prompt]
                        res = client.models.generate_content(model="gemini-3.6-flash", contents=contents).text.strip()
                        add_daily_report_to_word(doc, json.loads(res.replace("```json","").replace("```","").strip()), global_mohrir)
                elif audio_file:
                    res = client.models.generate_content(model="gemini-3.6-flash", contents=[audio_part, prompt]).text.strip()
                    add_daily_report_to_word(doc, json.loads(res.replace("```json","").replace("```","").strip()), global_mohrir)
                    
                bio = io.BytesIO()
                doc.save(bio)
                st.success("✅ Generated successfully!")
                st.download_button("⬇️ Download Daily Report", data=bio.getvalue(), file_name="Court_Daily_Report.docx")
            except Exception as e: st.error(f"Error: {e}")

# ----------------- UI: SOOCHNA -----------------
elif doc_type == "Soochna (सूचना - Jamanat / NBW Recall)":
    action_choice = st.radio("📌 Select Notice Type:", ["जमानत (Bail)", "NBW RECALL"])
    audio_file_soochna = st.audio_input("🎙️ Record Dictation")
    uploaded_files_soochna = st.file_uploader("Upload FIR/Photos (Optional)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    
    with st.expander("📝 Manual Text Overrides"):
        s_col1, s_col2 = st.columns(2)
        s_court = s_col1.text_input("न्यायालय (Court Name):")
        s_accused = s_col1.text_area("नाम पता अभियुक्त (Accused):")
        s_date = s_col1.text_input("उपस्थिति दिनांक (Date):")
        s_crime = s_col2.text_input("मु०अ०सं० (Crime No):")
        s_dhara = s_col2.text_input("धारा (Sections):")
        s_thana = s_col2.text_input("थाना (Thana):")

    if (uploaded_files_soochna or audio_file_soochna or s_court or s_accused) and st.button("Generate Soochna Document", type="primary"):
        with st.spinner("Processing..."):
            doc = Document()
            prompt = f"Extract details for 'Soochna' into JSON: court_name, accused, crime_no, dhara, thana, date. Overrides: Court:{s_court}, Accused:{s_accused}, CrimeNo:{s_crime}, Dhara:{s_dhara}, Thana:{s_thana}, Date:{s_date}. Output ONLY JSON."
            audio_part = types.Part.from_bytes(data=audio_file_soochna.getvalue(), mime_type=audio_file_soochna.type) if audio_file_soochna else None
            contents = []
            if uploaded_files_soochna: contents.extend([PIL.Image.open(f) for f in uploaded_files_soochna])
            if audio_part: contents.append(audio_part)
            contents.append(prompt)
            
            try:
                res = client.models.generate_content(model="gemini-3.6-flash", contents=contents).text.strip()
                add_soochna_to_word(doc, json.loads(res.replace("```json","").replace("```","").strip()), action_choice, global_mohrir)
                bio = io.BytesIO()
                doc.save(bio)
                st.success(f"✅ {action_choice} Soochna generated!")
                st.download_button("⬇️ Download Soochna.docx", data=bio.getvalue(), file_name="Soochna_Notice.docx")
            except Exception as e: st.error(f"Error: {e}")

# ----------------- UI: WITNESS REPORT -----------------
elif doc_type == "Witness/Summons Report (साक्षी/समन विवरण)":
    audio_file_wit = st.audio_input("🎙️ Record Dictation")
    uploaded_files_wit = st.file_uploader("Upload Table Photo (Optional)", type=["jpg", "jpeg", "png"])
    
    with st.expander("📝 Manual Number Entry"):
        w_court = st.text_input("न्यायालय (Court Name):", value="JM (JD) OUTLINE COURT कादीपुर सुल्तानपुर")
        w_date = st.text_input("दिनांक (Date):", value="25/07/2026")
        w_col1, w_col2, w_col3, w_col4 = st.columns(4)
        w_janpad = w_col1.text_input("जनपद:", value="सुल्तानपुर")
        w_1 = w_col2.text_input("साक्ष्य हेतु नियत वाद:", value="9")
        w_2 = w_col3.text_input("निर्गत कुल सम्मन:", value="9")
        w_3 = w_col4.text_input("तामीला होकर प्राप्त सम्मन:", value="8")
        w_col5, w_col6, w_col7, w_col8 = st.columns(4)
        w_4 = w_col5.text_input("निर्गत कुल वारंट:", value="00")
        w_5 = w_col6.text_input("तामीला होकर प्राप्त वारंट:", value="00")
        w_6 = w_col7.text_input("उपस्थित साक्षियों की संख्या:", value="06")
        w_7 = w_col8.text_input("परीक्षित साक्षियों की संख्या:", value="06")
        w_vc = st.text_input("VC Count:", value="0")

    if st.button("Generate Witness Report", type="primary"):
        doc = Document()
        data = {"court_name": w_court, "report_date": w_date, "janpad": w_janpad, "niyat_vaad": w_1, "nirgat_samman": w_2, "tamila_samman": w_3, "nirgat_warrant": w_4, "tamila_warrant": w_5, "upasthit_sakshi": w_6, "parikshit_sakshi": w_7, "vc_count": w_vc}
        add_witness_report_to_word(doc, data, global_mohrir)
        bio = io.BytesIO()
        doc.save(bio)
        st.success("✅ Generated!")
        st.download_button("⬇️ Download Witness Report.docx", data=bio.getvalue(), file_name="Witness_Report.docx")

# ----------------- UI: TOP 10 REPORT -----------------
elif doc_type == "Top 10 Report (टॉप 10 की सूचना)":
    st.info("📋 Edit the JSON table rows below or use the microphone to dictate updates.")
    audio_top10 = st.audio_input("🎙️ Record Dictation for Top 10 Report")
    default_top10 = [
        ["पंकज सिंह S/O राजेंद्र प्रताप सिंह", "464/16", "385 IPC", "कादीपुर Str", "3229/24", "4", "01", "03", "साक्ष्य में", "11-08-26"],
        ["-", "435/22", "3/25 A.act", "कादीपुर Str", "868/24", "6", "0", "6", "हाजिरी में", "11-08-26"],
        ["-", "109/15", "174A IPC", "कादीपुर Str", "4584/24", "11", "0", "11", "साक्ष्य में", "11-08-26"],
        ["-", "91/15", "386, 435, 427 IPC", "कादीपुर Str", "5451/24", "7", "0", "7", "साक्ष्य में", "11-08-26"]
    ]
    with st.expander("✏️ Manual Data Entry & Dictation (Edit Table Rows)", expanded=True):
        edited_text = st.text_area("Edit table rows (JSON format):", value=json.dumps(default_top10, ensure_ascii=False, indent=2), height=250)
        
    if st.button("Generate Top 10 Report", type="primary"):
        with st.spinner("Generating document..."):
            doc = Document()
            rows = json.loads(edited_text)
            if audio_top10:
                prompt = f"Given these current Top 10 table rows: {json.dumps(rows)}, update any dates, witness counts, or statuses based on this audio dictation. Return ONLY updated JSON list of lists."
                audio_part = types.Part.from_bytes(data=audio_top10.getvalue(), mime_type=audio_top10.type)
                res = client.models.generate_content(model="gemini-3.6-flash", contents=[audio_part, prompt]).text.strip()
                if "```json" in res: res = res.split("```json")[1].split("```")[0].strip()
                elif "```" in res: res = res.split("```")[1].split("```")[0].strip()
                rows = json.loads(res)
            add_top_10_report_to_word(doc, rows)
            bio = io.BytesIO()
            doc.save(bio)
            st.success("✅ Top 10 Report generated successfully in Landscape Mode!")
            st.download_button("⬇️ Download Top 10 Report.docx", data=bio.getvalue(), file_name="Top_10_Report.docx")

# ----------------- UI: ADVOCATE CASES REPORT -----------------
elif doc_type == "Cases Against Advocates Report (अधिवक्ता के विरुद्ध अभियोग)":
    st.info("⚖️ Edit the JSON table rows below or use the microphone to dictate status updates.")
    audio_adv = st.audio_input("🎙️ Record Dictation for Advocate Cases Report")
    default_adv = [
        ["35", "0415/16", "323/504/506 भादवि", "सुल्तानपुर", "कादीपुर", "झीगर पुत्र अनिकेत", "ACJM, कादीपुर", "हाजिरी", "18-08-26", "समन जारी है"],
        ["37", "474/2021", "323/504/506/427 भादवि", "सुल्तानपुर", "कादीपुर", "हरिराम पुत्र छंगूराम", "ACJM, कादीपुर", "हाजिरी", "21-09-26", "समन जारी है"],
        ["38", "423/11", "323/504/427 भादवि", "सुल्तानपुर", "कादीपुर", "दिवाकर सिंह पुत्र स्व० अनमोल सिंह", "ACJM, कादीपुर", "हाजिरी", "04-09-26", "समन जारी है"],
        ["39", "423/2011", "323/504/506 भादवि", "सुल्तानपुर", "कादीपुर", "सुरेश नारायण यादव पुत्र हरिंस यादव", "ACJM, कादीपुर", "हाजिरी", "04-09-26", "समन जारी है"],
        ["-", "402/21", "323,504,506 IPC", "सुल्तानपुर", "कादीपुर", "नितेश पांडेय s/o जय शंकर pandey", "ACJM KADIPUR", "हाजिरी", "22-09-26", "अभियुक्त को NBW"],
        ["-", "140/17", "323,506IPC", "सुल्तानपुर", "कादीपुर", "जिगर डूबे व अन्य", "ACJM KADIPUR", "हाजिरी", "16-08-26", "अभियुक्त को NBW"]
    ]
    with st.expander("✏️ Manual Data Entry & Dictation (Edit Table Rows)", expanded=True):
        edited_adv_text = st.text_area("Edit table rows (JSON format):", value=json.dumps(default_adv, ensure_ascii=False, indent=2), height=280)
        
    if st.button("Generate Advocate Cases Report", type="primary"):
        with st.spinner("Generating document..."):
            doc = Document()
            rows = json.loads(edited_adv_text)
            if audio_adv:
                prompt = f"Given these current Advocate cases rows: {json.dumps(rows)}, update any dates, actions, or statuses based on this audio dictation. Return ONLY updated JSON list of lists."
                audio_part = types.Part.from_bytes(data=audio_adv.getvalue(), mime_type=audio_adv.type)
                res = client.models.generate_content(model="gemini-3.6-flash", contents=[audio_part, prompt]).text.strip()
                if "```json" in res: res = res.split("```json")[1].split("```")[0].strip()
                elif "```" in res: res = res.split("```")[1].split("```")[0].strip()
                rows = json.loads(res)
            add_advocate_cases_report_to_word(doc, rows)
            bio = io.BytesIO()
            doc.save(bio)
            st.success("✅ Advocate Cases Report generated successfully in Landscape Mode!")
            st.download_button("⬇️ Download Advocate Cases Report.docx", data=bio.getvalue(), file_name="Advocate_Cases_Report.docx")

# ----------------- UI: REMAND SHEET (NEW) -----------------
elif doc_type == "Remand Sheet (रिमाण्ड शीट)":
    st.info("📄 Upload a police document/FIR or use voice dictation to auto-fill, or type manually below.")
    audio_remand = st.audio_input("🎙️ Record Dictation for Remand Sheet")
    uploaded_remand = st.file_uploader("Upload FIR / Case Photo", type=["jpg", "jpeg", "png"])
    
    with st.expander("📝 Manual Text Inputs & Overrides", expanded=True):
        r_col1, r_col2 = st.columns(2)
        r_court = r_col1.text_input("न्यायालय (Court Name):", value="अपर मुख्य न्यायिक मजिस्ट्रेट-SD कादीपुर सुलतानपुर")
        r_crime = r_col1.text_input("मु०अ०सं० (Crime No):")
        r_dhara = r_col1.text_input("धारा (Sections):")
        r_thana = r_col2.text_input("थाना (Thana):")
        r_janpad = r_col2.text_input("जनपद (District):", value="सुल्तानपुर")
        r_arrest = r_col2.text_input("गिरफ्तारी दिनांक (Arrest Date):")
        
        r_accused = st.text_input("अभियुक्त का नाम (Accused Name):")
        r_c1, r_c2, r_c3 = st.columns(3)
        r_from = r_c1.text_input("रिमाण्ड प्रारम्भ दिनांक (Remand From):")
        r_to = r_c2.text_input("रिमाण्ड समाप्ति दिनांक (Remand To):")
        r_jail = r_c3.text_input("जिला कारागार से तलब दिनांक (Jail Date):")
        
        r_po = st.text_input("पी.ओ. / न्यायाधीश हस्ताक्षर नाम (Judge Signature Box):", value="अपर मुख्य न्यायिक मजिस्ट्रेट-SD\nकादीपुर, सुलतानपुर।")
        r_date = st.text_input("दिनांक (Sign Date):")

    if (uploaded_remand or audio_remand or r_crime) and st.button("Generate Remand Sheet", type="primary"):
        with st.spinner("Processing Remand Sheet..."):
            doc = Document()
            prompt = f"""
            Extract details for a legal 'Remand Sheet' into JSON with EXACTLY these keys: "court_name", "crime_no", "dhara", "thana", "janpad", "arrest_date", "accused_name", "remand_from", "remand_to", "jail_date", "po_name", "date_sign".
            OVERRIDES (Use exactly if provided): Court:{r_court}, Crime:{r_crime}, Dhara:{r_dhara}, Thana:{r_thana}, Janpad:{r_janpad}, ArrestDate:{r_arrest}, Accused:{r_accused}, From:{r_from}, To:{r_to}, JailDate:{r_jail}, PO:{r_po}, Date:{r_date}. Output ONLY JSON.
            """
            audio_part = types.Part.from_bytes(data=audio_remand.getvalue(), mime_type=audio_remand.type) if audio_remand else None
            contents = []
            if uploaded_remand: contents.append(PIL.Image.open(uploaded_remand))
            if audio_part: contents.append(audio_part)
            contents.append(prompt)
            
            try:
                res = client.models.generate_content(model="gemini-3.6-flash", contents=contents).text.strip()
                if "```json" in res: res = res.split("```json")[1].split("```")[0].strip()
                elif "```" in res: res = res.split("```")[1].split("```")[0].strip()
                data = json.loads(res)
                
                add_remand_sheet_to_word(doc, data)
                bio = io.BytesIO()
                doc.save(bio)
                st.success("✅ Remand Sheet generated successfully!")
                st.download_button("⬇️ Download Remand_Sheet.docx", data=bio.getvalue(), file_name="Remand_Sheet.docx")
            except Exception as e: st.error(f"Error: {e}")

# ----------------- UI: COGNIZANCE / SANJAN ORDER (NEW) -----------------
elif doc_type == "Cognizance / Sanjan Order (संज्ञान आदेश)":
    st.info("⚖️ Upload a chargesheet photo or dictate details to generate the Sanjan Order.")
    audio_sanjan = st.audio_input("🎙️ Record Dictation for Sanjan Order")
    uploaded_sanjan = st.file_uploader("Upload Chargesheet Photo", type=["jpg", "jpeg", "png"])
    
    with st.expander("📝 Manual Text Inputs & Overrides", expanded=True):
        sn_col1, sn_col2 = st.columns(2)
        sn_court = sn_col1.text_input("न्यायालय (Court Name):", value="अपर मुख्य न्यायिक मजिस्ट्रेट, कादीपुर, सुलतानपुर")
        sn_upasthit = sn_col1.text_input("उपस्थित (Present):", value="सहायक अभियोजन अधिकारी / कोर्ट मोहर्रिर")
        sn_vaad = sn_col1.text_input("वाद सं० (Vaad No):")
        sn_comp = sn_col2.text_input("कम्प्यूटर फाइलिंग सं० (Computer Filing No):")
        sn_cr = sn_col2.text_input("सी०एन०आर०सं० (CNR No):")
        sn_cs = sn_col2.text_input("आरोप पत्र सं० (Chargesheet No):")
        
        sn_c1, sn_c2, sn_c3 = st.columns(3)
        sn_crime = sn_c1.text_input("मुकदमा अपराध संख्या (Crime No):")
        sn_dhara = sn_c2.text_input("अन्तर्गत धारा (Initial Sections):")
        sn_cog_dhara = sn_c3.text_input("संज्ञान धारा (Cognizance Sections):")
        
        sn_d1, sn_d2 = st.columns(2)
        sn_thana = sn_d1.text_input("थाना (Thana):")
        sn_janpad = sn_d1.text_input("जनपद (District):", value="सुल्तानपुर")
        sn_date = sn_d2.text_input("दिनांक (Order Date):")
        sn_next = sn_d2.text_input("वास्ते उपस्थिति दिनांक (Next Date):")
        
        sn_accused = st.text_area("अभियुक्तगण का नाम (Accused List):", value="1....................................................")
        sn_po = st.text_area("पी.ओ. / न्यायाधीश हस्ताक्षर बॉक्स (Judge Signature Box):", value="(विश्वजीत सिंह)\nअपर मुख्य न्यायिक मजिस्ट्रेट,\nकादीपुर जनपद सुल्तानपुर।")

    if (uploaded_sanjan or audio_sanjan or sn_crime) and st.button("Generate Sanjan Order", type="primary"):
        with st.spinner("Processing Sanjan Order..."):
            doc = Document()
            prompt = f"""
            Extract details for a legal 'Cognizance/Sanjan Order' into JSON with EXACTLY these keys: "court_name", "upasthit", "vaad_no", "comp_no", "cr_no", "chargesheet_no", "crime_no", "dhara", "thana", "janpad", "order_date", "accused_list", "initial_dhara", "cognizance_dhara", "next_date", "po_name".
            OVERRIDES (Use exactly if provided): Court:{sn_court}, Upasthit:{sn_upasthit}, VaadNo:{sn_vaad}, Comp:{sn_comp}, CR:{sn_cr}, CS:{sn_cs}, Crime:{sn_crime}, Dhara:{sn_dhara}, Thana:{sn_thana}, Janpad:{sn_janpad}, OrderDate:{sn_date}, Accused:{sn_accused}, InitialDhara:{sn_dhara}, CognizanceDhara:{sn_cog_dhara}, NextDate:{sn_next}, PO:{sn_po}. Output ONLY JSON.
            """
            audio_part = types.Part.from_bytes(data=audio_sanjan.getvalue(), mime_type=audio_sanjan.type) if audio_sanjan else None
            contents = []
            if uploaded_sanjan: contents.append(PIL.Image.open(uploaded_sanjan))
            if audio_part: contents.append(audio_part)
            contents.append(prompt)
            
            try:
                res = client.models.generate_content(model="gemini-3.6-flash", contents=contents).text.strip()
                if "```json" in res: res = res.split("```json")[1].split("```")[0].strip()
                elif "```" in res: res = res.split("```")[1].split("```")[0].strip()
                data = json.loads(res)
                
                add_sanjan_order_to_word(doc, data)
                bio = io.BytesIO()
                doc.save(bio)
                st.success("✅ Sanjan Order generated successfully!")
                st.download_button("⬇️ Download Sanjan_Order.docx", data=bio.getvalue(), file_name="Sanjan_Order.docx")
            except Exception as e: st.error(f"Error: {e}")
